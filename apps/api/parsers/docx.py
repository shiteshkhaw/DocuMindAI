import io
import asyncio
import hashlib
from typing import Dict, Any
import docx
from parsers.base import BaseParser, ParsedDocument, ParsedPage

class DocxParser(BaseParser):
    def can_handle(self, mime_type: str) -> bool:
        return mime_type.lower() in [
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "application/msword"
        ] or mime_type.lower().endswith("docx") or mime_type.lower().endswith("doc")

    async def parse(self, file_content: bytes, filename: str) -> ParsedDocument:
        return await asyncio.to_thread(self._parse_sync, file_content, filename)

    def _parse_sync(self, file_content: bytes, filename: str) -> ParsedDocument:
        stream = io.BytesIO(file_content)
        doc = docx.Document(stream)
        
        full_text_list = []
        for paragraph in doc.paragraphs:
            if paragraph.text:
                full_text_list.append(paragraph.text)
        
        full_text = "\n".join(full_text_list)
        
        # Docx has no native exact page count, so we treat it as 1 page or partition it
        pages = [ParsedPage(
            page_number=1,
            text=full_text,
            metadata={"source_page": 1}
        )]

        # Extract core properties if available
        props = doc.core_properties
        title = props.title if props.title else filename.rsplit(".", 1)[0]
        author = props.author if props.author else "Unknown"

        doc_metadata: Dict[str, Any] = {
            "title": title,
            "author": author,
            "page_count": 1,
            "file_size": len(file_content),
            "mime_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "checksum": hashlib.sha256(file_content).hexdigest(),
        }

        return ParsedDocument(pages=pages, metadata=doc_metadata)
